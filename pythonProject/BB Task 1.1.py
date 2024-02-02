import os
import platform
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import win32com.client

def capture_screenshot(url):
    # Set up Chrome options
    options = Options()
    options.add_argument("--window-size=1920,1080")
    # Initialize Chrome WebDriver with options
    driver = webdriver.Chrome(options=options)

    # Navigate to the specified URL
    driver.get(url)

    # Capture a screenshot
    screenshot_path = 'screenshot.png'
    driver.save_screenshot(screenshot_path)

    # Close the browser
    driver.quit()

def create_word_doc(file_path, image_path):
    document = Document()

    document.add_heading('Google Webpage Screenshot', 0).bold = True
    document.add_page_break()
    document.add_picture(image_path, width=Inches(6))

    # Add a header
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run('Google Screenshot Task')
    run.font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a footer
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Wesley 2024')
    run.font.size = Pt(9)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save(file_path)

def convert_to_pdf(docx_path, pdf_path):
    if platform.system() == "Windows":
        word_app = win32com.client.Dispatch("Word.Application")
        doc = word_app.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 represents PDF format
        doc.Close()
        word_app.Quit()
    else:
        raise OSError("Conversion to PDF is only supported on Windows.")

webpage = 'https://google.com'
capture_screenshot(webpage)

docx = 'C:/Users/user/Downloads/output.docx'
pdf = 'C:/Users/user/Downloads/output.pdf'
image = 'C:/Users/user/PycharmProjects/pythonProject/screenshot.png'

create_word_doc(docx, image)
convert_to_pdf(docx, pdf)


